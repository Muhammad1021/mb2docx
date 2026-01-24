"""DOCX Writer - V9 Forensic Verbatim Format.

This writer generates Word documents with exact visual parity to the
Bob_Frok_CV.docx and Bob_Frok_CL.docx exemplars:

- Names: 18pt bold centered
- Contact: 10pt centered (pipe-separated)
- Section headings: 12pt bold, ALL CAPS
- Job entries: PIPE separator "Title | Date" (NOT TABs)
- Institutions: 11pt italic, pipe for location
- Bullets: Text bullets for ATS compatibility
- All spacing values extracted forensically from exemplars
"""
from __future__ import annotations

import logging
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from .config import DocxStyleConfig
from .model import (
    AddressBlock,
    Block,
    ClosingBlock,
    ContactBlock,
    DateLineBlock,
    HeadingBlock,
    InstitutionBlock,
    JobEntryBlock,
    ListBlock,
    ParagraphBlock,
    SalutationBlock,
    SectionHeadingBlock,
)

log = logging.getLogger(__name__)


@dataclass(frozen=True)
class WriteResult:
    """Result of a write operation."""
    path: Path


# ============================================================================
# DOCUMENT SETUP
# ============================================================================

def _apply_page_setup(doc: Document, cfg: DocxStyleConfig) -> None:
    """Set page margins - FORENSIC: 1.0" all sides."""
    section = doc.sections[0]
    margin = Inches(cfg.margin_inches)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _apply_base_style(doc: Document, cfg: DocxStyleConfig) -> None:
    """Set default font for Normal style."""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = cfg.font_name
        normal.font.size = Pt(cfg.body_size_pt)
    except KeyError:
        pass


# ============================================================================
# TEXT HELPERS
# ============================================================================

def _add_run(paragraph, text: str, cfg: DocxStyleConfig,
             bold: bool = False, italic: bool = False,
             size_pt: Optional[int] = None) -> None:
    """Add a text run with explicit font settings.

    CRITICAL: Only set bold=True or italic=True, NEVER explicitly set False.
    Setting False writes <w:b w:val="0"/> which explicitly disables bold.
    """
    run = paragraph.add_run(text)
    run.font.name = cfg.font_name
    run.font.size = Pt(size_pt if size_pt else cfg.body_size_pt)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _parse_bold_markup(text: str) -> List[tuple]:
    """Parse **bold** markup into (text, is_bold) tuples."""
    parts = []
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0

    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()

    if last_end < len(text):
        parts.append((text[last_end:], False))

    if not parts:
        parts.append((text, False))

    return parts


# ============================================================================
# BLOCK RENDERERS
# ============================================================================

def _add_name(doc: Document, blk: HeadingBlock, cfg: DocxStyleConfig) -> None:
    """Render name: 18pt, BOLD, CENTER, ALL CAPS."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(cfg.space_after_name_pt)
    _add_run(p, blk.text.upper(), cfg, bold=True, size_pt=cfg.name_size_pt)


def _add_contact(doc: Document, blk: ContactBlock, cfg: DocxStyleConfig) -> None:
    """Render contact: 10pt, CENTER, pipe-separated."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(cfg.space_after_contact_pt)
    _add_run(p, blk.text, cfg, size_pt=cfg.contact_size_pt)


def _add_section_heading(doc: Document, blk: SectionHeadingBlock, cfg: DocxStyleConfig) -> None:
    """Render section heading: 12pt (FORENSIC), BOLD, ALL CAPS."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(cfg.space_before_section_pt)
    p.paragraph_format.space_after = Pt(cfg.space_after_section_pt)
    _add_run(p, blk.text.upper(), cfg, bold=True, size_pt=cfg.section_heading_size_pt)


def _add_job_entry(doc: Document, blk: JobEntryBlock, cfg: DocxStyleConfig) -> None:
    """Render job entry with DYNAMIC separator based on config.

    PIPE pattern (Bob Frok exemplar): "Senior Delivery Manager | June 2018 – Present"
    TAB pattern (traditional): "Senior Delivery Manager\t[right-aligned]June 2018 – Present"

    FORENSIC ANALYSIS CONFIRMED: Bob Frok uses PIPE, not TAB.
    V9.2 FIX: Only title is bold, date is NOT bold. Strip ** markers.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(cfg.space_after_job_entry_pt)

    # V9.2 FIX: Strip ** markdown markers from title
    title = blk.title.replace("**", "")

    if cfg.job_entry_separator == "PIPE":
        # ═══════════════════════════════════════════════════════════════════
        # PIPE SEPARATOR (Bob Frok style) - THE CORRECT WAY
        # V9.2: Title bold, separator and date NOT bold
        # ═══════════════════════════════════════════════════════════════════
        _add_run(p, title, cfg, bold=True)
        _add_run(p, f" | {blk.date_range}", cfg)

    elif cfg.job_entry_separator == "TAB":
        # TAB SEPARATOR (traditional style - NOT used by Bob Frok)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(cfg.tab_stop_inches), WD_TAB_ALIGNMENT.RIGHT)
        _add_run(p, title, cfg, bold=True)
        p.add_run('\t')
        _add_run(p, blk.date_range, cfg)

    else:
        # NONE - just the title
        _add_run(p, title, cfg, bold=True)


def _add_institution(doc: Document, blk: InstitutionBlock, cfg: DocxStyleConfig) -> None:
    """Render institution: 11pt, ITALIC."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(cfg.space_after_institution_pt)
    _add_run(p, blk.text, cfg, italic=True)


def _add_paragraph(doc: Document, blk: ParagraphBlock, cfg: DocxStyleConfig) -> None:
    """Render paragraph with inline **bold** support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(cfg.space_after_paragraph_pt)

    parts = _parse_bold_markup(blk.text)
    for text, is_bold in parts:
        _add_run(p, text, cfg, bold=is_bold)


def _add_list(doc: Document, blk: ListBlock, cfg: DocxStyleConfig,
              is_last_before_section: bool = False) -> None:
    """Render list items as bullet paragraphs with proper hanging indent.

    V9.2 FIX: Text after bullet aligns with wrapped lines at 0.5".
    - Bullet at 0.25" (first line pulled back by hanging indent)
    - Tab stop at 0.5" aligns text
    - Wrapped lines start at 0.5" (left_indent)
    """
    num_items = len(blk.items)
    for idx, item in enumerate(blk.items):
        p = doc.add_paragraph()

        # Gold standard indent: left=0.5", hanging=0.25"
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)

        # Tab stop at 0.5" so text aligns with wrapped lines
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))

        is_last = (idx == num_items - 1)

        if is_last and is_last_before_section:
            p.paragraph_format.space_after = Pt(cfg.space_after_last_bullet_pt)
        elif is_last:
            p.paragraph_format.space_after = Pt(cfg.space_after_last_bullet_pt)
        else:
            p.paragraph_format.space_after = Pt(cfg.space_after_bullet_pt)

        # Text bullet followed by tab for alignment (ATS compatible)
        bullet = "•\t" if not blk.ordered else f"{idx + 1}.\t"
        _add_run(p, bullet, cfg)

        parts = _parse_bold_markup(item)
        for text, is_bold in parts:
            _add_run(p, text, cfg, bold=is_bold)


def _add_date_line(doc: Document, blk: DateLineBlock, cfg: DocxStyleConfig) -> None:
    """Render cover letter date line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    _add_run(p, blk.text, cfg)


def _add_address_block(doc: Document, blk: AddressBlock, cfg: DocxStyleConfig) -> None:
    """Render address block."""
    for idx, line in enumerate(blk.lines):
        p = doc.add_paragraph()
        is_last = (idx == len(blk.lines) - 1)
        p.paragraph_format.space_after = Pt(20 if is_last else 0)
        _add_run(p, line, cfg)


def _add_salutation(doc: Document, blk: SalutationBlock, cfg: DocxStyleConfig) -> None:
    """Render salutation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _add_run(p, blk.text, cfg)


def _add_closing(doc: Document, blk: ClosingBlock, cfg: DocxStyleConfig) -> None:
    """Render closing, signature, phone, and email on separate lines."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    _add_run(p1, blk.closing, cfg)

    if blk.signature:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        _add_run(p2, blk.signature, cfg)

    if blk.phone:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        _add_run(p3, blk.phone, cfg)

    if blk.email:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(0)
        _add_run(p4, blk.email, cfg)


# ============================================================================
# MAIN RENDERING
# ============================================================================

def render_blocks(doc: Document, blocks: Iterable[Block], cfg: DocxStyleConfig) -> None:
    """Render all blocks to document with context-aware spacing."""
    blocks_list = list(blocks)
    num_blocks = len(blocks_list)

    for i, blk in enumerate(blocks_list):
        next_blk = blocks_list[i + 1] if i + 1 < num_blocks else None
        is_last_before_section = isinstance(next_blk, (SectionHeadingBlock, JobEntryBlock))

        if isinstance(blk, HeadingBlock):
            if blk.level == 1:
                _add_name(doc, blk, cfg)
            else:
                p = doc.add_paragraph()
                _add_run(p, blk.text, cfg, bold=True)
        elif isinstance(blk, ContactBlock):
            _add_contact(doc, blk, cfg)
        elif isinstance(blk, SectionHeadingBlock):
            _add_section_heading(doc, blk, cfg)
        elif isinstance(blk, JobEntryBlock):
            _add_job_entry(doc, blk, cfg)
        elif isinstance(blk, InstitutionBlock):
            _add_institution(doc, blk, cfg)
        elif isinstance(blk, ListBlock):
            _add_list(doc, blk, cfg, is_last_before_section)
        elif isinstance(blk, ParagraphBlock):
            _add_paragraph(doc, blk, cfg)
        elif isinstance(blk, DateLineBlock):
            _add_date_line(doc, blk, cfg)
        elif isinstance(blk, AddressBlock):
            _add_address_block(doc, blk, cfg)
        elif isinstance(blk, SalutationBlock):
            _add_salutation(doc, blk, cfg)
        elif isinstance(blk, ClosingBlock):
            _add_closing(doc, blk, cfg)
        else:
            log.warning("Unknown block type: %s", type(blk))


def new_document(*, cfg: DocxStyleConfig, title: Optional[str] = None,
                 author: Optional[str] = None) -> Document:
    """Create a new document with base styling."""
    doc = Document()
    _apply_page_setup(doc, cfg)
    _apply_base_style(doc, cfg)
    if title or author:
        cp = doc.core_properties
        if title:
            cp.title = title
        if author:
            cp.author = author
    return doc


def write_document(blocks: Iterable[Block], output_path: Path,
                   cfg: DocxStyleConfig, author: Optional[str] = None) -> WriteResult:
    """Write blocks to a DOCX file."""
    doc = new_document(cfg=cfg, author=author)
    render_blocks(doc, blocks, cfg)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    log.info("Wrote document: %s", output_path)
    return WriteResult(path=output_path)


def build_docx(blocks: Iterable[Block], *, cfg: DocxStyleConfig,
               title: Optional[str] = None, author: Optional[str] = None) -> Document:
    """Build a complete document from blocks."""
    doc = new_document(cfg=cfg, title=title, author=author)
    render_blocks(doc, blocks, cfg)
    return doc


def safe_save_docx(doc: Document, out_path: Path) -> WriteResult:
    """Write to temp file then replace (atomic save).

    Falls back to direct save if atomic replace fails (e.g., file open in Word).
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(out_path.parent)) as tf:
        tmp_path = Path(tf.name)

    try:
        doc.save(str(tmp_path))
        try:
            tmp_path.replace(out_path)
            log.info("Wrote %s", out_path)
            return WriteResult(path=out_path)
        except PermissionError:
            # Target file might be open - try direct save instead
            tmp_path.unlink()  # Clean up temp file
            try:
                doc.save(str(out_path))
                log.info("Wrote %s (direct save)", out_path)
                return WriteResult(path=out_path)
            except PermissionError as e:
                # File is definitely locked - give helpful error
                raise PermissionError(
                    f"Cannot save to '{out_path.name}' - the file may be open in Word or another program. "
                    f"Please close it and try again."
                ) from e
    finally:
        if tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass


def add_page_break(doc: Document) -> None:
    """Add a page break."""
    doc.add_page_break()
